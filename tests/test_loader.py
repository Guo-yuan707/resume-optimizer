"""loader 模块的测试(第 15 课新增的文件读取函数)。

测什么:
    - extract_text() 能把不同格式的"内存字节"抽成纯文本:
      .txt(utf-8 / gbk 兜底)、.md、.docx(段落 + 表格)、.pdf(文字型)
    - 不支持的格式 / 没有可提取文字的扫描版 PDF → 抛友好的 ValueError
    - 图片 OCR 太重、太慢,不进单测,靠实现时手工冒烟验证

怎么跑(在项目根目录):
    .venv/Scripts/python -m pytest tests/test_loader.py -v
"""
from io import BytesIO

import pytest

from resume_optimizer.loader import extract_text


# ---------- .txt / .md:纯文本 ----------

def test_txt_utf8():
    """utf-8 编码的中文 txt:应该原样读出。"""
    data = "专业技能\nPython 数据分析".encode("utf-8")
    assert extract_text("简历.txt", data) == "专业技能\nPython 数据分析"


def test_txt_gbk_fallback():
    """GBK/GB18030 编码的中文 txt(Windows 记事本常见):utf-8 解不了,
    应该自动回退到 gb18030 兜底解出中文。"""
    data = "专业技能\nPython".encode("gbk")
    text = extract_text("简历.txt", data)
    assert "专业技能" in text and "Python" in text


def test_md():
    """markdown 本质是纯文本,后缀 .md 也要能读。"""
    data = "# JD\n要求 Python + RAG".encode("utf-8")
    assert "Python + RAG" in extract_text("jd.md", data)


# ---------- .docx ----------

def test_docx_paragraph_and_table():
    """docx 里正文放在段落和表格两种载体里,都应该被读到。"""
    import docx

    doc = docx.Document()
    doc.add_paragraph("专业技能")
    doc.add_paragraph("Python 数据分析")
    table = doc.add_table(rows=2, cols=1)
    table.rows[0].cells[0].text = "教育经历"
    table.rows[1].cells[0].text = "郑州财经学院 本科"

    buf = BytesIO()
    doc.save(buf)
    text = extract_text("简历.docx", buf.getvalue())

    assert "专业技能" in text
    assert "Python 数据分析" in text
    assert "教育经历" in text
    assert "郑州财经学院" in text  # 表格里的字没被漏掉


# ---------- .pdf ----------

def _make_mini_pdf(text: str) -> bytes:
    """手写一个最小的、含一行 ASCII 文字的合法 PDF。

    为什么自己造?pypdf 只负责"读",不负责"写"文字型 PDF;
    造一个最小 PDF 作为测试夹具,证明 pdf 读取链路是通的。
    中文 PDF 需要内嵌字体,手写不了 —— 留到真实文件手工验证。
    """
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
    ]
    stream = b"BT /F1 24 Tf 72 700 Td (" + text.encode("ascii") + b") Tj ET"
    objects.append(b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream")
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]  # offsets[i] = 第 i 个对象在文件里的偏移(xref 表要用)
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += (f"{i} 0 obj\n").encode() + obj + b"\nendobj\n"

    xref_pos = len(out)
    out += b"xref\n0 " + str(len(objects) + 1).encode() + b"\n"
    out += b"0000000000 65535 f \n"          # 第 0 个对象:free 对象(固定)
    for off in offsets[1:]:
        out += (f"{off:010d} 00000 n \n").encode()
    out += (b"trailer\n<< /Size " + str(len(objects) + 1).encode()
            + b" /Root 1 0 R >>\nstartxref\n" + str(xref_pos).encode() + b"\n%%EOF")
    return bytes(out)


def test_pdf_text():
    """文字型 PDF:能抽出那一行文字。"""
    text = extract_text("简历.pdf", _make_mini_pdf("Hello Resume Python"))
    assert "Hello Resume Python" in text


def test_pdf_no_text_raises():
    """没有可提取文字的 PDF(相当于扫描版/图片型):应抛友好 ValueError。"""
    with pytest.raises(ValueError):
        extract_text("scan.pdf", _make_mini_pdf(""))


# ---------- 兜底与报错 ----------

def test_unknown_extension_raises():
    """不支持的格式:抛 ValueError,并且提示里列出支持的格式。"""
    with pytest.raises(ValueError) as excinfo:
        extract_text("简历.xlsx", b"whatever")
    assert "不支持" in str(excinfo.value)
    assert ".pdf" in str(excinfo.value)


def test_no_extension_raises():
    """连扩展名都没有的文件:同样该报错,而不是悄悄当 txt。"""
    with pytest.raises(ValueError):
        extract_text("无扩展名文件", b"content")
